"""
Generate Model Documentation (DOCX) for Test Datasets
Creates comprehensive SR 11-7 compliant documentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_application_high_documentation():
    """Create documentation for Set 4: Application Scorecard (High Performance)"""
    print("Creating documentation for Set 4: Application Scorecard (High Performance)...")

    doc = Document()

    # Title
    title = doc.add_heading('Application Scorecard Model Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('High Performance Credit Origination Model')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph(f'Document Version: 1.0')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph(f'Model Type: Application Scorecard')
    doc.add_paragraph(f'Use Case: Credit Origination / New Customer Acquisition')

    # 1. Model Purpose and Use Cases
    doc.add_heading('1. Model Purpose and Use Cases', level=1)
    doc.add_paragraph(
        'This application scorecard is designed to assess credit risk for new customer acquisitions. '
        'The model predicts the probability of default within 12 months of account origination.'
    )
    doc.add_paragraph('Primary Use Cases:')
    doc.add_paragraph('• Credit decisioning for new applicants', style='List Bullet')
    doc.add_paragraph('• Credit line assignment', style='List Bullet')
    doc.add_paragraph('• Risk-based pricing', style='List Bullet')
    doc.add_paragraph('• Portfolio risk management', style='List Bullet')

    # 2. Conceptual Soundness
    doc.add_heading('2. Conceptual Soundness', level=1)
    doc.add_paragraph('Theoretical Foundation:')
    doc.add_paragraph(
        'The model is based on established credit risk theory and logistic regression methodology. '
        'It incorporates key risk drivers identified in academic literature and industry best practices.'
    )
    doc.add_paragraph('Key Assumptions:')
    doc.add_paragraph('• Historical patterns are indicative of future behavior', style='List Bullet')
    doc.add_paragraph('• Features are independent and additive in their risk contribution', style='List Bullet')
    doc.add_paragraph('• Population characteristics remain stable over time', style='List Bullet')
    doc.add_paragraph('• Economic conditions remain within historical ranges', style='List Bullet')

    # 3. Data Quality and Representativeness
    doc.add_heading('3. Data Quality and Representativeness', level=1)
    doc.add_paragraph('Development Sample:')
    doc.add_paragraph('• Training: 2,000 accounts', style='List Bullet')
    doc.add_paragraph('• Test: 1,000 accounts', style='List Bullet')
    doc.add_paragraph('• Out-of-Time: 600 accounts', style='List Bullet')
    doc.add_paragraph('• Default Rate: 7% (within expected range)', style='List Bullet')

    doc.add_paragraph('Data Quality Metrics:')
    doc.add_paragraph('• Completeness: 100% (no missing values)', style='List Bullet')
    doc.add_paragraph('• Accuracy: Verified through multiple sources', style='List Bullet')
    doc.add_paragraph('• Representativeness: Covers full credit spectrum', style='List Bullet')

    # 4. Model Performance
    doc.add_heading('4. Model Performance', level=1)
    doc.add_paragraph('Discriminatory Power:')
    doc.add_paragraph('• KS Statistic: > 0.35 (Excellent)', style='List Bullet')
    doc.add_paragraph('• Gini Coefficient: > 0.45 (Excellent)', style='List Bullet')
    doc.add_paragraph('• AUC-ROC: > 0.75 (Strong)', style='List Bullet')

    doc.add_paragraph('Calibration:')
    doc.add_paragraph('• Hosmer-Lemeshow Test: Passed', style='List Bullet')
    doc.add_paragraph('• Predicted vs Actual: Well-aligned', style='List Bullet')

    # 5. Stability Analysis
    doc.add_heading('5. Stability Analysis', level=1)
    doc.add_paragraph('Population Stability:')
    doc.add_paragraph('• PSI (Train vs Test): < 0.10 (Very Stable)', style='List Bullet')
    doc.add_paragraph('• PSI (Train vs OOT): < 0.10 (Very Stable)', style='List Bullet')
    doc.add_paragraph('• CSI: < 0.15 (Stable)', style='List Bullet')

    doc.add_paragraph('Temporal Stability:')
    doc.add_paragraph('• Performance consistent across time periods', style='List Bullet')
    doc.add_paragraph('• No significant drift detected', style='List Bullet')

    # 6. Assumptions Testing
    doc.add_heading('6. Assumptions Testing', level=1)
    doc.add_paragraph('All key assumptions have been tested and validated:')
    doc.add_paragraph('• Linearity: Confirmed through residual analysis', style='List Bullet')
    doc.add_paragraph('• Independence: VIF < 5 for all features', style='List Bullet')
    doc.add_paragraph('• Stability: Confirmed through time-series analysis', style='List Bullet')

    doc.add_paragraph('Sensitivity Analysis:')
    doc.add_paragraph('Model performance remains robust under:')
    doc.add_paragraph('• ±10% change in default rate', style='List Bullet')
    doc.add_paragraph('• ±20% change in feature distributions', style='List Bullet')
    doc.add_paragraph('• Economic stress scenarios', style='List Bullet')

    # 7. Implementation and Deployment
    doc.add_heading('7. Implementation and Deployment', level=1)
    doc.add_paragraph('Implementation Details:')
    doc.add_paragraph('• Platform: Python-based scoring engine', style='List Bullet')
    doc.add_paragraph('• Integration: Real-time API', style='List Bullet')
    doc.add_paragraph('• Latency: < 100ms per score', style='List Bullet')
    doc.add_paragraph('• Availability: 99.9% uptime', style='List Bullet')

    doc.add_paragraph('Validation:')
    doc.add_paragraph('• Code review: Completed', style='List Bullet')
    doc.add_paragraph('• UAT: Passed', style='List Bullet')
    doc.add_paragraph('• Production testing: Successful', style='List Bullet')

    # 8. Ongoing Monitoring
    doc.add_heading('8. Ongoing Monitoring', level=1)
    doc.add_paragraph('Monitoring Plan:')
    doc.add_paragraph('• Performance tracking: Monthly', style='List Bullet')
    doc.add_paragraph('• PSI monitoring: Monthly', style='List Bullet')
    doc.add_paragraph('• Drift detection: Automated alerts', style='List Bullet')
    doc.add_paragraph('• Full revalidation: Annually', style='List Bullet')

    doc.add_paragraph('Key Metrics Monitored:')
    doc.add_paragraph('• Default rate', style='List Bullet')
    doc.add_paragraph('• KS statistic', style='List Bullet')
    doc.add_paragraph('• Population stability (PSI)', style='List Bullet')
    doc.add_paragraph('• Feature distributions', style='List Bullet')

    # 9. Documentation and Governance
    doc.add_heading('9. Documentation and Governance', level=1)
    doc.add_paragraph('Documentation Maintained:')
    doc.add_paragraph('• Model development documentation', style='List Bullet')
    doc.add_paragraph('• Validation reports', style='List Bullet')
    doc.add_paragraph('• Technical specifications', style='List Bullet')
    doc.add_paragraph('• User guides', style='List Bullet')
    doc.add_paragraph('• Change logs', style='List Bullet')

    doc.add_paragraph('Governance:')
    doc.add_paragraph('• Model Risk Committee oversight', style='List Bullet')
    doc.add_paragraph('• Independent validation: Annual', style='List Bullet')
    doc.add_paragraph('• Regulatory reporting: As required', style='List Bullet')

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This application scorecard demonstrates excellent predictive power, strong stability, '
        'and comprehensive documentation. It meets all SR 11-7 requirements and is suitable for '
        'production deployment in credit origination processes.'
    )

    # Save
    doc.save('set4_application_high/application_scorecard_documentation.docx')
    print("✅ Set 4 documentation created: set4_application_high/application_scorecard_documentation.docx\n")


def create_collections_early_documentation():
    """Create documentation for Set 5: Collections Early Stage"""
    print("Creating documentation for Set 5: Collections Early Stage...")

    doc = Document()

    # Title
    title = doc.add_heading('Collections Early Stage Scorecard Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Early Delinquency Management Model (30-90 DPD)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(128, 0, 0)

    doc.add_paragraph(f'Document Version: 1.0')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph(f'Model Type: Collections Scorecard - Early Stage')
    doc.add_paragraph(f'Use Case: Early Delinquency Management (30-90 Days Past Due)')

    # 1. Model Purpose and Use Cases
    doc.add_heading('1. Model Purpose and Use Cases', level=1)
    doc.add_paragraph(
        'This collections scorecard predicts the probability of account recovery for customers '
        'in early-stage delinquency (30-90 days past due). The model helps optimize collection '
        'strategies and resource allocation.'
    )
    doc.add_paragraph('Primary Use Cases:')
    doc.add_paragraph('• Recovery probability assessment', style='List Bullet')
    doc.add_paragraph('• Collection strategy optimization', style='List Bullet')
    doc.add_paragraph('• Resource allocation and prioritization', style='List Bullet')
    doc.add_paragraph('• Contact strategy determination', style='List Bullet')
    doc.add_paragraph('• Early intervention targeting', style='List Bullet')

    # 2. Conceptual Soundness
    doc.add_heading('2. Conceptual Soundness', level=1)
    doc.add_paragraph('Theoretical Foundation:')
    doc.add_paragraph(
        'The model is based on collections best practices and behavioral economics principles. '
        'It incorporates factors known to influence recovery likelihood in early delinquency.'
    )
    doc.add_paragraph('Key Assumptions:')
    doc.add_paragraph('• Early intervention improves recovery rates', style='List Bullet')
    doc.add_paragraph('• Customer behavior patterns are predictive', style='List Bullet')
    doc.add_paragraph('• Account characteristics influence recovery', style='List Bullet')
    doc.add_paragraph('• Contact effectiveness varies by customer segment', style='List Bullet')

    # 3. Data Quality and Representativeness
    doc.add_heading('3. Data Quality and Representativeness', level=1)
    doc.add_paragraph('Development Sample:')
    doc.add_paragraph('• Training: 1,500 accounts', style='List Bullet')
    doc.add_paragraph('• Test: 800 accounts', style='List Bullet')
    doc.add_paragraph('• Out-of-Time: 500 accounts', style='List Bullet')
    doc.add_paragraph('• Recovery Rate: 45% (typical for early stage)', style='List Bullet')

    doc.add_paragraph('Data Quality Metrics:')
    doc.add_paragraph('• Completeness: 100%', style='List Bullet')
    doc.add_paragraph('• Accuracy: Verified through collections system', style='List Bullet')
    doc.add_paragraph('• Representativeness: Covers 30-90 DPD spectrum', style='List Bullet')

    # 4. Model Performance
    doc.add_heading('4. Model Performance', level=1)
    doc.add_paragraph('Discriminatory Power:')
    doc.add_paragraph('• KS Statistic: 0.20-0.30 (Good for collections)', style='List Bullet')
    doc.add_paragraph('• Gini Coefficient: 0.25-0.35 (Acceptable)', style='List Bullet')
    doc.add_paragraph('• AUC-ROC: 0.65-0.70 (Adequate)', style='List Bullet')

    doc.add_paragraph('Note: Lower thresholds are expected for collections models due to:')
    doc.add_paragraph('• Higher inherent uncertainty in delinquent populations', style='List Bullet')
    doc.add_paragraph('• External factors beyond model scope', style='List Bullet')
    doc.add_paragraph('• Behavioral volatility in stressed customers', style='List Bullet')

    # 5. Stability Analysis
    doc.add_heading('5. Stability Analysis', level=1)
    doc.add_paragraph('Population Stability:')
    doc.add_paragraph('• PSI (Train vs Test): < 0.25 (Stable)', style='List Bullet')
    doc.add_paragraph('• PSI (Train vs OOT): < 0.25 (Stable)', style='List Bullet')
    doc.add_paragraph('• CSI: < 0.25 (Acceptable)', style='List Bullet')

    doc.add_paragraph('Note: Higher PSI tolerance for collections due to:')
    doc.add_paragraph('• More volatile delinquent population', style='List Bullet')
    doc.add_paragraph('• Seasonal variations in delinquency', style='List Bullet')
    doc.add_paragraph('• Economic sensitivity', style='List Bullet')

    # 6. Assumptions Testing
    doc.add_heading('6. Assumptions Testing', level=1)
    doc.add_paragraph('All key assumptions have been tested:')
    doc.add_paragraph('• Recovery patterns: Validated through historical data', style='List Bullet')
    doc.add_paragraph('• Feature importance: Confirmed through SHAP analysis', style='List Bullet')
    doc.add_paragraph('• Temporal stability: Tested across multiple periods', style='List Bullet')

    doc.add_paragraph('Sensitivity Analysis:')
    doc.add_paragraph('Model remains effective under:')
    doc.add_paragraph('• ±15% change in recovery rate', style='List Bullet')
    doc.add_paragraph('• Varying economic conditions', style='List Bullet')
    doc.add_paragraph('• Different collection strategies', style='List Bullet')

    # 7. Implementation and Deployment
    doc.add_heading('7. Implementation and Deployment', level=1)
    doc.add_paragraph('Implementation Details:')
    doc.add_paragraph('• Platform: Collections management system', style='List Bullet')
    doc.add_paragraph('• Integration: Batch and real-time scoring', style='List Bullet')
    doc.add_paragraph('• Update frequency: Daily', style='List Bullet')
    doc.add_paragraph('• Decision support: Automated recommendations', style='List Bullet')

    doc.add_paragraph('Validation:')
    doc.add_paragraph('• System integration: Tested', style='List Bullet')
    doc.add_paragraph('• Collections team training: Completed', style='List Bullet')
    doc.add_paragraph('• Pilot program: Successful', style='List Bullet')

    # 8. Ongoing Monitoring
    doc.add_heading('8. Ongoing Monitoring', level=1)
    doc.add_paragraph('Monitoring Plan:')
    doc.add_paragraph('• Recovery rate tracking: Weekly', style='List Bullet')
    doc.add_paragraph('• Model performance: Monthly', style='List Bullet')
    doc.add_paragraph('• PSI monitoring: Monthly', style='List Bullet')
    doc.add_paragraph('• Strategy effectiveness: Quarterly', style='List Bullet')
    doc.add_paragraph('• Full revalidation: Annually', style='List Bullet')

    doc.add_paragraph('Key Metrics Monitored:')
    doc.add_paragraph('• Recovery rate by score band', style='List Bullet')
    doc.add_paragraph('• Contact effectiveness', style='List Bullet')
    doc.add_paragraph('• Time to recovery', style='List Bullet')
    doc.add_paragraph('• Population stability', style='List Bullet')

    # 9. Documentation and Governance
    doc.add_heading('9. Documentation and Governance', level=1)
    doc.add_paragraph('Documentation Maintained:')
    doc.add_paragraph('• Model development documentation', style='List Bullet')
    doc.add_paragraph('• Validation reports', style='List Bullet')
    doc.add_paragraph('• Collections strategy guides', style='List Bullet')
    doc.add_paragraph('• Performance monitoring reports', style='List Bullet')
    doc.add_paragraph('• Regulatory compliance documentation', style='List Bullet')

    doc.add_paragraph('Governance:')
    doc.add_paragraph('• Collections Committee oversight', style='List Bullet')
    doc.add_paragraph('• Model Risk Committee review', style='List Bullet')
    doc.add_paragraph('• Independent validation: Annual', style='List Bullet')
    doc.add_paragraph('• FDCPA compliance: Ongoing', style='List Bullet')

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This collections early stage scorecard provides effective recovery probability predictions '
        'for accounts in early delinquency. It meets SR 11-7 requirements adapted for collections '
        'models and supports optimized collection strategies and resource allocation.'
    )

    # Save
    doc.save('set5_collections_early/collections_early_documentation.docx')
    print("✅ Set 5 documentation created: set5_collections_early/collections_early_documentation.docx\n")


if __name__ == "__main__":
    print("="*80)
    print("Generating Model Documentation (DOCX)")
    print("="*80 + "\n")

    create_application_high_documentation()
    create_collections_early_documentation()

    print("="*80)
    print("✅ All documentation generated successfully!")
    print("="*80)
    print("\nGenerated:")
    print("  • set4_application_high/application_scorecard_documentation.docx")
    print("  • set5_collections_early/collections_early_documentation.docx")

# Made with Bob
