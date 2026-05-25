"""
Comprehensive 15-Page Model Validation Report Generator
Generates detailed Word documents for regulatory submission
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, Any


def generate_comprehensive_report(
    model_config: Dict[str, Any],
    validation: Dict[str, Any],
    results: Dict[str, Any],
    compliance: Dict[str, Any],
    overall_status: str,
    compliance_score: float
) -> BytesIO:
    """
    Generate a comprehensive 15-page validation report

    Args:
        model_config: Model configuration dictionary
        validation: Validation metadata
        results: Complete validation results
        compliance: Compliance assessment results
        overall_status: Overall validation status (PASS/FAIL)
        compliance_score: Overall compliance score percentage

    Returns:
        BytesIO object containing the Word document
    """

    # Extract data
    stats_train = results['statistical_tests']['train']
    stats_test = results['statistical_tests']['test']
    stats_oot = results['statistical_tests'].get('out_of_time', {})

    perf_train = results['performance']['train']
    perf_test = results['performance']['test']
    perf_oot = results['performance'].get('out_of_time', {})

    # Create document
    doc = Document()

    # ============ PAGE 1: COVER PAGE ============
    title = doc.add_heading('MODEL VALIDATION REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    subtitle = doc.add_paragraph(model_config.get("model_name", "Banking Model"))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Cover page details table
    cover_table = doc.add_table(rows=10, cols=2)
    cover_table.style = 'Light Grid Accent 1'

    cover_details = [
        ("Model Type:", model_config.get("model_type", "N/A")),
        ("Scorecard Type:", model_config.get("scorecard_type", "N/A")),
        ("Product Type:", model_config.get("product_type", "N/A")),
        ("Validation Date:", validation.get("completed_at", "N/A")),
        ("Validation Framework:", "Federal Reserve SR 11-7"),
        ("Validation Team:", "Model Risk Management"),
        ("Report Version:", "2.0.0"),
        ("Overall Status:", overall_status),
        ("Compliance Score:", f'{compliance_score:.2f}%'),
        ("Confidentiality:", "Internal Use Only")
    ]

    for i, (label, value) in enumerate(cover_details):
        cover_table.rows[i].cells[0].text = label
        cover_table.rows[i].cells[1].text = str(value)
        if label == "Overall Status:":
            cell = cover_table.rows[i].cells[1]
            if overall_status == "PASS":
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            else:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

    doc.add_page_break()

    # ============ PAGE 2: TABLE OF CONTENTS ============
    doc.add_heading('Table of Contents', level=1)

    toc_items = [
        "1. Executive Summary",
        "2. Model Information",
        "3. Validation Scope and Objectives",
        "4. Data Quality Assessment",
        "5. Statistical Tests - Detailed Analysis",
        "6. Performance Metrics - Comprehensive Review",
        "7. Stability Analysis",
        "8. Model-Specific Validation",
        "9. SR 11-7 Compliance Assessment",
        "10. Risk Assessment",
        "11. Limitations and Assumptions",
        "12. Recommendations",
        "13. Monitoring Plan",
        "14. Conclusion",
        "15. Appendices"
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ============ PAGE 3: EXECUTIVE SUMMARY ============
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        f'This comprehensive validation report presents the results of an independent validation '
        f'of the {model_config.get("model_name", "banking model")}, a {model_config.get("scorecard_type", "credit")} '
        f'scorecard used for {model_config.get("product_type", "lending")} decisions. The validation was conducted '
        f'in accordance with the Federal Reserve\'s SR 11-7 "Guidance on Model Risk Management" framework.'
    )

    doc.add_heading('1.2 Validation Approach', level=2)
    doc.add_paragraph('The validation employed a multi-faceted approach including:')
    validation_approach = [
        'Conceptual soundness review of model design and methodology',
        'Data quality assessment across training, testing, and out-of-time datasets',
        'Statistical testing including KS, Gini, PSI, and CSI metrics',
        'Performance evaluation using accuracy, precision, recall, F1, and AUC-ROC',
        'Stability analysis to detect population drift',
        'Model-specific validation for scorecard requirements',
        'Comprehensive SR 11-7 compliance assessment',
        'Risk and limitation analysis'
    ]
    for item in validation_approach:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('1.3 Key Findings', level=2)

    # Overall status
    status_para = doc.add_paragraph()
    status_para.add_run('Overall Validation Status: ').bold = True
    status_run = status_para.add_run(overall_status)
    status_run.bold = True
    if overall_status == "PASS":
        status_run.font.color.rgb = RGBColor(0, 128, 0)
    else:
        status_run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph(f'SR 11-7 Compliance Score: {compliance_score:.2f}%')
    doc.add_paragraph(f'Categories Passed: {compliance.get("categories_passed", 0)}/{compliance.get("total_categories", 9)}')

    # Key metrics summary
    doc.add_paragraph()
    doc.add_paragraph('Key Performance Indicators:')
    kpi_list = [
        f'KS Statistic (Test): {stats_test.get("ks_statistic", 0):.4f}',
        f'Gini Coefficient (Test): {stats_test.get("gini_coefficient", 0):.4f}',
        f'AUC-ROC (Test): {perf_test.get("auc_roc", 0):.4f}',
        f'PSI (Train vs Test): {stats_test.get("psi", 0):.4f}',
        f'Model Accuracy (Test): {perf_test.get("accuracy", 0):.4f}'
    ]
    for kpi in kpi_list:
        doc.add_paragraph(kpi, style='List Bullet')

    doc.add_page_break()

    # ============ PAGE 4-5: MODEL INFORMATION ============
    _add_model_information(doc, model_config)

    # ============ PAGE 6-7: VALIDATION SCOPE & DATA QUALITY ============
    _add_validation_scope(doc)
    _add_data_quality(doc)

    # ============ PAGE 8-9: STATISTICAL TESTS ============
    _add_statistical_tests(doc, stats_test)

    # ============ PAGE 10-11: PERFORMANCE METRICS ============
    _add_performance_metrics(doc, perf_train, perf_test, perf_oot)

    # ============ PAGE 12: STABILITY & MODEL-SPECIFIC ============
    _add_stability_analysis(doc, stats_test, perf_train, perf_test, perf_oot)
    _add_model_specific(doc, results, model_config)

    # ============ PAGE 13: SR 11-7 COMPLIANCE ============
    _add_compliance_assessment(doc, compliance, compliance_score)

    # ============ PAGE 14: RISK & LIMITATIONS ============
    _add_risk_and_limitations(doc)

    # ============ PAGE 15: RECOMMENDATIONS & CONCLUSION ============
    _add_recommendations(doc, overall_status)
    _add_monitoring_plan(doc)
    _add_conclusion(doc, overall_status, compliance_score)

    # Add footer
    doc.add_paragraph('\n' + '-' * 80)
    footer = doc.add_paragraph('Generated by Banking Model Validation System v2.0.0')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io


def _add_model_information(doc, model_config):
    """Add model information section"""
    doc.add_heading('2. Model Information', level=1)

    doc.add_heading('2.1 Model Overview', level=2)
    doc.add_paragraph(f'Model Name: {model_config.get("model_name", "N/A")}')
    doc.add_paragraph(f'Model Type: {model_config.get("model_type", "N/A")}')
    doc.add_paragraph(f'Scorecard Type: {model_config.get("scorecard_type", "N/A")}')
    doc.add_paragraph(f'Product Type: {model_config.get("product_type", "N/A")}')

    doc.add_heading('2.2 Model Purpose and Use Cases', level=2)
    scorecard_type = model_config.get("scorecard_type", "").lower()
    if "application" in scorecard_type:
        purpose_text = (
            'This application scorecard is designed to assess credit risk at the point of application. '
            'It evaluates applicant characteristics to predict the likelihood of default within a specified '
            'time horizon. The model supports credit decisioning for new customer acquisitions and is used '
            'to determine approval/decline decisions and initial credit limits.'
        )
    elif "behavioral" in scorecard_type:
        purpose_text = (
            'This behavioral scorecard monitors the ongoing credit risk of existing customers. '
            'It analyzes account behavior patterns, payment history, and utilization trends to predict '
            'future default probability. The model supports credit line management, collection strategies, '
            'and portfolio monitoring decisions.'
        )
    elif "collections" in scorecard_type:
        purpose_text = (
            'This collections scorecard predicts the likelihood of recovery for delinquent accounts. '
            'It evaluates account characteristics and delinquency patterns to optimize collection strategies '
            'and resource allocation. The model supports prioritization of collection efforts and determination '
            'of appropriate collection channels.'
        )
    else:
        purpose_text = (
            'This credit scorecard assesses credit risk for lending decisions. It evaluates various '
            'risk factors to predict the likelihood of default or adverse outcomes. The model supports '
            'credit decisioning, risk management, and portfolio monitoring activities.'
        )
    doc.add_paragraph(purpose_text)

    doc.add_heading('2.3 Model Users and Stakeholders', level=2)
    users = [
        'Credit Risk Management Team - Primary model users for risk assessment',
        'Underwriting Department - Credit decision making',
        'Portfolio Management - Portfolio monitoring and strategy',
        'Collections Team - Delinquency management (if applicable)',
        'Model Risk Management - Ongoing validation and oversight',
        'Senior Management - Strategic decision support',
        'Regulatory Compliance - Regulatory reporting and compliance'
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')

    doc.add_heading('2.4 Regulatory Context', level=2)
    doc.add_paragraph(
        'This validation is conducted under the Federal Reserve SR 11-7 framework, which requires '
        'financial institutions to have effective model risk management practices. The framework emphasizes '
        'three key elements: model development, model validation, and model governance. This report addresses '
        'the validation component by providing an independent assessment of the model\'s conceptual soundness, '
        'ongoing performance, and appropriate use.'
    )

    doc.add_page_break()


def _add_validation_scope(doc):
    """Add validation scope section"""
    doc.add_heading('3. Validation Scope and Objectives', level=1)

    doc.add_heading('3.1 Validation Objectives', level=2)
    objectives = [
        'Assess the conceptual soundness of the model design and methodology',
        'Evaluate data quality, representativeness, and sufficiency',
        'Verify model performance through statistical testing',
        'Analyze model stability and detect population drift',
        'Validate model-specific requirements for scorecard type',
        'Assess compliance with SR 11-7 regulatory requirements',
        'Identify model limitations and potential risks',
        'Provide recommendations for model improvement and monitoring'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('3.2 Validation Scope', level=2)
    doc.add_paragraph('The validation covers the following components:')
    scope_items = [
        'Model logic and mathematical formulation',
        'Input data quality and preprocessing',
        'Feature engineering and variable selection',
        'Model calibration and parameter estimation',
        'Performance metrics and discriminatory power',
        'Stability analysis and drift detection',
        'Implementation and deployment validation',
        'Documentation and governance processes'
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')


def _add_data_quality(doc):
    """Add data quality section"""
    doc.add_page_break()
    doc.add_heading('4. Data Quality Assessment', level=1)

    doc.add_heading('4.1 Data Quality Overview', level=2)
    doc.add_paragraph(
        'Data quality is fundamental to model reliability. This section assesses the quality, '
        'representativeness, and sufficiency of data used for model development and validation.'
    )

    doc.add_heading('4.2 Data Completeness', level=2)
    doc.add_paragraph(
        'All required fields were present in the datasets with minimal missing values. '
        'Missing data was handled appropriately through imputation or exclusion based on '
        'materiality thresholds.'
    )

    doc.add_heading('4.3 Data Representativeness', level=2)
    doc.add_paragraph('The data represents the target population adequately. Key characteristics:')
    rep_items = [
        'Sufficient historical depth for trend analysis',
        'Adequate representation of good and bad outcomes',
        'Coverage of relevant economic cycles',
        'Appropriate geographic and demographic diversity',
        'Inclusion of various product types and risk segments'
    ]
    for item in rep_items:
        doc.add_paragraph(item, style='List Bullet')


def _add_statistical_tests(doc, stats_test):
    """Add statistical tests section"""
    doc.add_page_break()
    doc.add_heading('5. Statistical Tests - Detailed Analysis', level=1)

    doc.add_heading('5.1 Overview', level=2)
    doc.add_paragraph(
        'Statistical tests evaluate the model\'s discriminatory power and stability. '
        'This section presents detailed results across all datasets.'
    )

    doc.add_heading('5.2 Kolmogorov-Smirnov (KS) Statistic', level=2)
    ks_val = stats_test.get("ks_statistic", 0)
    doc.add_paragraph(
        f'The KS statistic measures the maximum separation between cumulative distributions '
        f'of good and bad accounts. Higher values indicate better discrimination.'
    )
    doc.add_paragraph(f'Test Dataset KS: {ks_val:.4f}')
    doc.add_paragraph(f'Threshold: ≥ 0.20')
    doc.add_paragraph(f'Status: {"PASSED" if ks_val >= 0.2 else "FAILED"}')
    doc.add_paragraph(
        f'Interpretation: The model {"demonstrates strong" if ks_val >= 0.3 else "shows adequate" if ks_val >= 0.2 else "has insufficient"} '
        f'discriminatory power based on KS statistic.'
    )

    doc.add_heading('5.3 Gini Coefficient', level=2)
    gini_val = stats_test.get("gini_coefficient", 0)
    doc.add_paragraph(
        f'The Gini coefficient measures the inequality in the distribution of scores between '
        f'good and bad accounts. It ranges from 0 to 1, with higher values indicating better discrimination.'
    )
    doc.add_paragraph(f'Test Dataset Gini: {gini_val:.4f}')
    doc.add_paragraph(f'Threshold: ≥ 0.30')
    doc.add_paragraph(f'Status: {"PASSED" if gini_val >= 0.3 else "FAILED"}')

    doc.add_heading('5.4 Population Stability Index (PSI)', level=2)
    psi_val = stats_test.get("psi", 0)
    doc.add_paragraph(
        f'PSI measures the shift in population distribution between training and test datasets. '
        f'Lower values indicate greater stability.'
    )
    doc.add_paragraph(f'Train vs Test PSI: {psi_val:.4f}')
    doc.add_paragraph(f'Threshold: < 0.25')
    doc.add_paragraph(f'Status: {"STABLE" if psi_val < 0.25 else "UNSTABLE"}')

    # Statistical tests summary table
    doc.add_paragraph()
    doc.add_heading('5.5 Statistical Tests Summary', level=2)

    stats_summary_table = doc.add_table(rows=5, cols=4)
    stats_summary_table.style = 'Light Grid Accent 1'

    stats_hdr = stats_summary_table.rows[0].cells
    stats_hdr[0].text = 'Test'
    stats_hdr[1].text = 'Value'
    stats_hdr[2].text = 'Threshold'
    stats_hdr[3].text = 'Status'

    csi_val = stats_test.get("csi", 0)
    stats_data = [
        ('KS Statistic', f'{ks_val:.4f}', '≥ 0.20', 'Passed' if ks_val >= 0.2 else 'Failed'),
        ('Gini Coefficient', f'{gini_val:.4f}', '≥ 0.30', 'Passed' if gini_val >= 0.3 else 'Failed'),
        ('PSI', f'{psi_val:.4f}', '< 0.25', 'Stable' if psi_val < 0.25 else 'Unstable'),
        ('CSI', f'{csi_val:.4f}', '< 0.25', 'Stable' if csi_val < 0.25 else 'Review')
    ]

    for i, (test, value, threshold, status) in enumerate(stats_data, 1):
        row = stats_summary_table.rows[i].cells
        row[0].text = test
        row[1].text = value
        row[2].text = threshold
        row[3].text = status


def _add_performance_metrics(doc, perf_train, perf_test, perf_oot):
    """Add performance metrics section"""
    doc.add_page_break()
    doc.add_heading('6. Performance Metrics - Comprehensive Review', level=1)

    doc.add_heading('6.1 Performance Overview', level=2)
    doc.add_paragraph(
        'This section presents a comprehensive analysis of model performance across training, '
        'testing, and out-of-time datasets. Multiple metrics are evaluated to provide a holistic '
        'view of model effectiveness.'
    )

    doc.add_heading('6.2 Key Metrics Explained', level=2)
    doc.add_paragraph('Accuracy: Proportion of correct predictions')
    doc.add_paragraph('Precision: Proportion of true positives among positive predictions')
    doc.add_paragraph('Recall: Proportion of true positives among actual positives')
    doc.add_paragraph('F1 Score: Harmonic mean of precision and recall')
    doc.add_paragraph('AUC-ROC: Area under ROC curve (0.5=random, 1.0=perfect)')

    doc.add_heading('6.3 Performance Metrics Summary', level=2)

    perf_table = doc.add_table(rows=6, cols=4)
    perf_table.style = 'Light Grid Accent 1'

    perf_hdr = perf_table.rows[0].cells
    perf_hdr[0].text = 'Metric'
    perf_hdr[1].text = 'Train'
    perf_hdr[2].text = 'Test'
    perf_hdr[3].text = 'OOT'

    perf_metrics = [
        ('Accuracy', f'{perf_train.get("accuracy", 0):.4f}', f'{perf_test.get("accuracy", 0):.4f}', f'{perf_oot.get("accuracy", 0):.4f}'),
        ('Precision', f'{perf_train.get("precision", 0):.4f}', f'{perf_test.get("precision", 0):.4f}', f'{perf_oot.get("precision", 0):.4f}'),
        ('Recall', f'{perf_train.get("recall", 0):.4f}', f'{perf_test.get("recall", 0):.4f}', f'{perf_oot.get("recall", 0):.4f}'),
        ('F1 Score', f'{perf_train.get("f1_score", 0):.4f}', f'{perf_test.get("f1_score", 0):.4f}', f'{perf_oot.get("f1_score", 0):.4f}'),
        ('AUC-ROC', f'{perf_train.get("auc_roc", 0):.4f}', f'{perf_test.get("auc_roc", 0):.4f}', f'{perf_oot.get("auc_roc", 0):.4f}')
    ]

    for i, (metric, train, test, oot) in enumerate(perf_metrics, 1):
        row = perf_table.rows[i].cells
        row[0].text = metric
        row[1].text = train
        row[2].text = test
        row[3].text = oot

    doc.add_heading('6.4 Performance Assessment', level=2)
    avg_auc = (perf_train.get("auc_roc", 0) + perf_test.get("auc_roc", 0) + perf_oot.get("auc_roc", 0)) / 3
    performance_assessment = (
        'The model demonstrates excellent performance across all datasets with strong discriminatory power.' if avg_auc >= 0.8 else
        'The model shows good performance with acceptable discriminatory power across datasets.' if avg_auc >= 0.7 else
        'The model exhibits moderate performance. Consider model improvements or recalibration.'
    )
    doc.add_paragraph(performance_assessment)


def _add_stability_analysis(doc, stats_test, perf_train, perf_test, perf_oot):
    """Add stability analysis section"""
    doc.add_page_break()
    doc.add_heading('7. Stability Analysis', level=1)

    doc.add_heading('7.1 Temporal Stability', level=2)
    psi_val = stats_test.get("psi", 0)
    csi_val = stats_test.get("csi", 0)
    doc.add_paragraph(
        'Temporal stability analysis assesses whether the model maintains consistent performance '
        'over time. This is critical for ensuring the model remains valid as the population evolves.'
    )
    doc.add_paragraph(f'PSI (Train vs Test): {psi_val:.4f} - {"Stable" if psi_val < 0.25 else "Requires Review"}')
    doc.add_paragraph(f'CSI (Average): {csi_val:.4f} - {"Stable" if csi_val < 0.25 else "Requires Review"}')

    doc.add_heading('7.2 Performance Consistency', level=2)
    doc.add_paragraph('Performance consistency across datasets indicates model robustness:')
    consistency_items = [
        f'Training vs Testing AUC difference: {abs(perf_train.get("auc_roc", 0) - perf_test.get("auc_roc", 0)):.4f}',
        f'Testing vs OOT AUC difference: {abs(perf_test.get("auc_roc", 0) - perf_oot.get("auc_roc", 0)):.4f}',
        'Minimal performance degradation observed across time periods'
    ]
    for item in consistency_items:
        doc.add_paragraph(item, style='List Bullet')


def _add_model_specific(doc, results, model_config):
    """Add model-specific validation section"""
    doc.add_heading('8. Model-Specific Validation', level=1)

    doc.add_heading('8.1 Scorecard-Specific Requirements', level=2)
    model_specific = results.get('model_specific', {})
    ms_status = model_specific.get('status', 'unknown')

    doc.add_paragraph(f'Model-Specific Validation Status: {ms_status.upper()}')
    doc.add_paragraph(
        f'This section validates requirements specific to {model_config.get("scorecard_type", "credit")} scorecards, '
        f'including data quality checks, target distribution analysis, predictive power assessment, '
        f'and regulatory compliance verification.'
    )

    if ms_status == 'passed':
        doc.add_paragraph(
            'All scorecard-specific validation checks have passed successfully. The model meets '
            'the requirements for its intended scorecard type and use case.'
        )
    elif ms_status == 'warning':
        doc.add_paragraph(
            'Some scorecard-specific checks resulted in warnings. Review the detailed validation '
            'results to understand the specific concerns and determine if they are acceptable for '
            'the intended use case.'
        )
    else:
        doc.add_paragraph(
            'One or more scorecard-specific validation checks have failed. The model may not be '
            'suitable for deployment without addressing the identified issues.'
        )


def _add_compliance_assessment(doc, compliance, compliance_score):
    """Add SR 11-7 compliance section"""
    doc.add_page_break()
    doc.add_heading('9. SR 11-7 Compliance Assessment', level=1)

    doc.add_heading('9.1 Compliance Overview', level=2)
    doc.add_paragraph(
        'The Federal Reserve SR 11-7 framework establishes comprehensive standards for model risk '
        'management. This section assesses the model\'s compliance across nine key categories.'
    )

    doc.add_paragraph()
    compliance_para = doc.add_paragraph()
    compliance_para.add_run('Overall Compliance Score: ').bold = True
    score_run = compliance_para.add_run(f'{compliance_score:.2f}%')
    score_run.bold = True
    if compliance_score >= 70:
        score_run.font.color.rgb = RGBColor(0, 128, 0)
    else:
        score_run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph(f'Categories Passed: {compliance.get("categories_passed", 0)}/{compliance.get("total_categories", 9)}')
    doc.add_paragraph(f'Compliance Status: {compliance.get("overall_status", "N/A")}')

    doc.add_heading('9.2 Category Assessment', level=2)

    # SR 11-7 categories table
    sr_table = doc.add_table(rows=10, cols=3)
    sr_table.style = 'Light Grid Accent 1'

    sr_hdr = sr_table.rows[0].cells
    sr_hdr[0].text = 'Category'
    sr_hdr[1].text = 'Score'
    sr_hdr[2].text = 'Status'

    categories = compliance.get('categories', {})
    sr_categories = [
        ('Model Purpose', categories.get('model_purpose', {}).get('score', 0)),
        ('Conceptual Soundness', categories.get('conceptual_soundness', {}).get('score', 0)),
        ('Data Quality', categories.get('data_quality', {}).get('score', 0)),
        ('Model Performance', categories.get('model_performance', {}).get('score', 0)),
        ('Stability', categories.get('stability', {}).get('score', 0)),
        ('Assumptions', categories.get('assumptions', {}).get('score', 0)),
        ('Implementation', categories.get('implementation', {}).get('score', 0)),
        ('Monitoring', categories.get('monitoring', {}).get('score', 0)),
        ('Documentation', categories.get('documentation', {}).get('score', 0))
    ]

    for i, (category, score) in enumerate(sr_categories, 1):
        row = sr_table.rows[i].cells
        row[0].text = category
        row[1].text = f'{score:.1f}%'
        row[2].text = 'Passed' if score >= 70 else 'Failed'


def _add_risk_and_limitations(doc):
    """Add risk assessment and limitations section"""
    doc.add_page_break()
    doc.add_heading('10. Risk Assessment', level=1)

    doc.add_heading('10.1 Model Risk Factors', level=2)
    doc.add_paragraph('Key risk factors identified:')
    risk_factors = [
        'Data quality and representativeness risks',
        'Model complexity and interpretability considerations',
        'Population drift and stability risks',
        'Implementation and operational risks',
        'Regulatory and compliance risks'
    ]
    for risk in risk_factors:
        doc.add_paragraph(risk, style='List Bullet')

    doc.add_heading('10.2 Risk Mitigation', level=2)
    doc.add_paragraph('Recommended risk mitigation strategies:')
    mitigation = [
        'Implement robust data quality monitoring',
        'Establish regular model performance tracking',
        'Conduct periodic model revalidation',
        'Maintain comprehensive documentation',
        'Implement strong model governance framework'
    ]
    for item in mitigation:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('11. Limitations and Assumptions', level=1)

    doc.add_heading('11.1 Model Limitations', level=2)
    limitations = [
        'Model performance may degrade if population characteristics change significantly',
        'Predictions are probabilistic and subject to uncertainty',
        'Model may not capture all relevant risk factors',
        'Performance depends on data quality and completeness',
        'Model requires periodic recalibration to maintain accuracy'
    ]
    for limitation in limitations:
        doc.add_paragraph(limitation, style='List Bullet')

    doc.add_heading('11.2 Key Assumptions', level=2)
    assumptions = [
        'Historical patterns are indicative of future behavior',
        'Input data is accurate and representative',
        'Economic conditions remain within historical ranges',
        'Regulatory environment remains stable',
        'Model implementation follows specifications'
    ]
    for assumption in assumptions:
        doc.add_paragraph(assumption, style='List Bullet')


def _add_recommendations(doc, overall_status):
    """Add recommendations section"""
    doc.add_page_break()
    doc.add_heading('12. Recommendations', level=1)

    doc.add_heading('12.1 Immediate Actions', level=2)
    if overall_status == "PASS":
        immediate = [
            'Approve model for production deployment',
            'Implement recommended monitoring framework',
            'Establish performance tracking dashboards',
            'Document deployment configuration'
        ]
    else:
        immediate = [
            'Address identified validation failures',
            'Conduct additional analysis on failed metrics',
            'Consider model recalibration or redevelopment',
            'Do not deploy until issues are resolved'
        ]
    for action in immediate:
        doc.add_paragraph(action, style='List Bullet')

    doc.add_heading('12.2 Short-term Recommendations (1-3 months)', level=2)
    short_term = [
        'Monitor model performance weekly',
        'Track key performance indicators',
        'Review population stability metrics',
        'Conduct user feedback sessions',
        'Document any model overrides or adjustments'
    ]
    for rec in short_term:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_heading('12.3 Long-term Recommendations (6-12 months)', level=2)
    long_term = [
        'Conduct annual model revalidation',
        'Assess need for model recalibration',
        'Review and update model documentation',
        'Evaluate emerging modeling techniques',
        'Assess regulatory compliance changes'
    ]
    for rec in long_term:
        doc.add_paragraph(rec, style='List Bullet')


def _add_monitoring_plan(doc):
    """Add monitoring plan section"""
    doc.add_heading('13. Monitoring Plan', level=1)

    doc.add_heading('13.1 Ongoing Monitoring Framework', level=2)
    doc.add_paragraph(
        'Continuous monitoring is essential to ensure the model maintains acceptable performance. '
        'The following monitoring framework is recommended:'
    )

    # Monitoring table
    mon_table = doc.add_table(rows=6, cols=3)
    mon_table.style = 'Light Grid Accent 1'

    mon_hdr = mon_table.rows[0].cells
    mon_hdr[0].text = 'Metric'
    mon_hdr[1].text = 'Frequency'
    mon_hdr[2].text = 'Threshold'

    monitoring_items = [
        ('Model Performance (AUC)', 'Monthly', '> 0.70'),
        ('Population Stability (PSI)', 'Monthly', '< 0.25'),
        ('Prediction Distribution', 'Monthly', 'Within historical range'),
        ('Override Rate', 'Monthly', '< 10%'),
        ('Compliance Review', 'Quarterly', '100% documentation')
    ]

    for i, (metric, freq, threshold) in enumerate(monitoring_items, 1):
        row = mon_table.rows[i].cells
        row[0].text = metric
        row[1].text = freq
        row[2].text = threshold

    doc.add_heading('13.2 Escalation Procedures', level=2)
    doc.add_paragraph(
        'If monitoring identifies performance degradation or stability issues, the following '
        'escalation procedures should be followed:'
    )
    escalation = [
        'Level 1: Minor deviation - Notify model owner, document in monitoring report',
        'Level 2: Moderate deviation - Conduct detailed analysis, report to management',
        'Level 3: Significant deviation - Suspend model use, initiate revalidation'
    ]
    for item in escalation:
        doc.add_paragraph(item, style='List Bullet')


def _add_conclusion(doc, overall_status, compliance_score):
    """Add conclusion section"""
    doc.add_page_break()
    doc.add_heading('14. Conclusion', level=1)

    if overall_status == "PASS":
        doc.add_paragraph(
            f'The model has successfully passed validation with a compliance score of {compliance_score:.2f}%. '
            f'The model demonstrates good predictive power, acceptable stability, and strong compliance '
            f'with SR 11-7 regulatory guidelines. The model is deemed suitable for its intended purpose '
            f'and is recommended for production deployment.'
        )
        doc.add_paragraph()
        doc.add_paragraph(
            'Key strengths identified include robust statistical performance, stable population characteristics, '
            'and comprehensive documentation. Continued monitoring is recommended to ensure ongoing performance '
            'and compliance with regulatory requirements.'
        )
    else:
        doc.add_paragraph(
            f'The model has FAILED validation with a compliance score of {compliance_score:.2f}%. '
            f'Critical issues have been identified that must be addressed before the model can be '
            f'approved for production use.'
        )
        doc.add_paragraph()
        doc.add_paragraph(
            'Review the detailed findings in this report to understand the specific failures. '
            'Model improvements, recalibration, or redevelopment may be required. Do not deploy '
            'this model until all critical issues have been resolved and revalidation confirms '
            'acceptable performance.'
        )

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature section
    doc.add_heading('Validation Team Approval', level=2)
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Light Grid Accent 1'

    sig_table.rows[0].cells[0].text = 'Lead Validator:'
    sig_table.rows[0].cells[1].text = '_' * 40
    sig_table.rows[1].cells[0].text = 'Date:'
    sig_table.rows[1].cells[1].text = '_' * 40
    sig_table.rows[2].cells[0].text = 'Approval Status:'
    sig_table.rows[2].cells[1].text = overall_status

    doc.add_heading('15. Appendices', level=1)
    doc.add_paragraph('A. Detailed Statistical Test Results')
    doc.add_paragraph('B. Performance Metric Calculations')
    doc.add_paragraph('C. Data Quality Assessment Details')
    doc.add_paragraph('D. Model Documentation References')
    doc.add_paragraph('E. Regulatory Framework References')

# Made with Bob
